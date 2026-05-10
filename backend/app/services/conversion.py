import os
import csv
from pathlib import Path
from fastapi import HTTPException


class ConversionService:
    """Handles all document conversion logic."""

    async def convert(self, input_path: str, source_format: str, target_format: str) -> str:
        output_path = os.path.splitext(input_path)[0] + f"_out.{target_format}"

        handlers = {
            ("pdf", "docx"): self._pdf_to_docx,
            ("jpg", "png"): self._image_to_image,
            ("png", "jpg"): self._image_to_image,
            ("jpg", "pdf"): self._image_to_pdf,
            ("png", "pdf"): self._image_to_pdf,
            ("xlsx", "csv"): self._xlsx_to_csv,
            ("xlsx", "pdf"): self._xlsx_to_pdf,
            ("xlsx", "docx"): self._xlsx_to_docx,
            ("docx", "pdf"): self._docx_to_pdf,
            ("docx", "xlsx"): self._docx_to_xlsx,
        }

        handler = handlers.get((source_format, target_format))
        if not handler:
            raise HTTPException(
                status_code=400,
                detail=f"No handler for {source_format} → {target_format}",
            )

        await handler(input_path, output_path, target_format)
        return output_path

    # ------------------------------------------------------------------ #\n    # PDF → Word (DOCX)  — text extraction via pypdf + python-docx
    # ------------------------------------------------------------------ #
    async def _pdf_to_docx(self, input_path: str, output_path: str, _target_format: str):
        try:
            import pypdf
            from docx import Document
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        try:
            reader = pypdf.PdfReader(input_path)
            doc = Document()
            for page in reader.pages:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    doc.add_paragraph(line)
                doc.add_paragraph()  # blank line between pages
            doc.save(output_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF to DOCX conversion failed: {e}")

    # ------------------------------------------------------------------ #
    # Image ↔ Image
    # ------------------------------------------------------------------ #
    async def _image_to_image(self, input_path: str, output_path: str, target_format: str):
        try:
            from PIL import Image
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        img = Image.open(input_path)
        if target_format == "jpg":
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(output_path, "JPEG", quality=95)
        else:
            img.save(output_path, "PNG")

    # ------------------------------------------------------------------ #
    # Image → PDF
    # ------------------------------------------------------------------ #
    async def _image_to_pdf(self, input_path: str, output_path: str, _target_format: str):
        try:
            from PIL import Image
            from reportlab.lib.pagesizes import letter  # type: ignore
            from reportlab.pdfgen import canvas as rl_canvas  # type: ignore
        except ImportError:
            # Fallback using only Pillow
            try:
                from PIL import Image

                img = Image.open(input_path)
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                img.save(output_path, "PDF", resolution=100.0)
                return
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Image to PDF failed: {e}")

        img = Image.open(input_path)
        img_width, img_height = img.size
        page_width, page_height = letter

        # Scale to fit letter page
        ratio = min(page_width / img_width, page_height / img_height)
        draw_width = img_width * ratio
        draw_height = img_height * ratio
        x = (page_width - draw_width) / 2
        y = (page_height - draw_height) / 2

        c = rl_canvas.Canvas(output_path, pagesize=letter)
        c.drawImage(input_path, x, y, width=draw_width, height=draw_height)
        c.save()

    # ------------------------------------------------------------------ #
    # Excel → CSV
    # ------------------------------------------------------------------ #
    async def _xlsx_to_csv(self, input_path: str, output_path: str, _target_format: str):
        try:
            import openpyxl
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
        ws = wb.active
        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            for row in ws.iter_rows(values_only=True):  # type: ignore
                writer.writerow([("" if v is None else v) for v in row])
        wb.close()

    # ------------------------------------------------------------------ #
    # Excel → PDF
    # ------------------------------------------------------------------ #
    async def _xlsx_to_pdf(self, input_path: str, output_path: str, _target_format: str):
        try:
            import openpyxl
            from reportlab.lib.pagesizes import letter, landscape  # type: ignore
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle  # type: ignore
            from reportlab.lib import colors  # type: ignore
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
        ws = wb.active
        data = []
        for row in ws.iter_rows(values_only=True):  # type: ignore
            data.append([str(v) if v is not None else "" for v in row])
        wb.close()

        if not data:
            raise HTTPException(status_code=400, detail="Excel file is empty")

        doc = SimpleDocTemplate(output_path, pagesize=landscape(letter))
        table = Table(data)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.lightblue]),
                ]
            )
        )
        doc.build([table])

    # ------------------------------------------------------------------ #
    # Excel → Word (DOCX)
    # ------------------------------------------------------------------ #
    async def _xlsx_to_docx(self, input_path: str, output_path: str, _target_format: str):
        try:
            import openpyxl
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        wb = openpyxl.load_workbook(input_path, read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))  # type: ignore
        wb.close()

        if not rows:
            raise HTTPException(status_code=400, detail="Excel file is empty")

        document = Document()
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, value in enumerate(row):
                table.cell(r_idx, c_idx).text = str(value) if value is not None else ""
        document.save(output_path)

    # ------------------------------------------------------------------ #
    # Word → PDF
    # ------------------------------------------------------------------ #
    async def _docx_to_pdf(self, input_path: str, output_path: str, _target_format: str):
        # Try LibreOffice first (best quality)
        import subprocess
        import shutil

        if shutil.which("libreoffice") or shutil.which("soffice"):
            cmd = shutil.which("libreoffice") or shutil.which("soffice")
            out_dir = os.path.dirname(output_path)
            result = subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf", "--outdir", out_dir, input_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0:
                # LibreOffice names the file after the input file
                lo_output = os.path.join(
                    out_dir,
                    os.path.splitext(os.path.basename(input_path))[0] + ".pdf",
                )
                if os.path.exists(lo_output) and lo_output != output_path:
                    os.rename(lo_output, output_path)
                return

        # Fallback: python-docx + reportlab
        try:
            from docx import Document  # type: ignore
            from reportlab.lib.pagesizes import letter  # type: ignore
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
            from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
            from reportlab.lib.units import inch  # type: ignore
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        doc = Document(input_path)
        pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        for para in doc.paragraphs:
            if para.text.strip():
                style = styles["Normal"]
                if para.style.name.startswith("Heading"):
                    style = styles.get("Heading1", styles["Normal"])
                story.append(Paragraph(para.text, style))
                story.append(Spacer(1, 0.1 * inch))

        if not story:
            story.append(Paragraph("(empty document)", styles["Normal"]))

        pdf_doc.build(story)

    # ------------------------------------------------------------------ #
    # Word → Excel (XLSX)
    # ------------------------------------------------------------------ #
    async def _docx_to_xlsx(self, input_path: str, output_path: str, _target_format: str):
        try:
            from docx import Document
            import openpyxl
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        doc = Document(input_path)
        wb = openpyxl.Workbook()
        ws = wb.active

        row_idx = 1
        # Extract tables first
        for table in doc.tables:
            for row in table.rows:
                ws.append([cell.text for cell in row.cells])
                row_idx += 1
            ws.append([])  # blank separator row between tables
            row_idx += 1

        # If no tables, fall back to writing paragraphs as rows
        if row_idx == 1:
            for para in doc.paragraphs:
                if para.text.strip():
                    ws.append([para.text])

        wb.save(output_path)

    # ------------------------------------------------------------------ #
    # Merge PDFs
    # ------------------------------------------------------------------ #
    async def merge_pdfs(self, input_paths: list) -> str:
        try:
            import pypdf
        except ImportError as e:
            raise HTTPException(status_code=500, detail=f"Missing dependency: {e}")

        import tempfile, uuid
        output_path = os.path.join(tempfile.gettempdir(), f"{uuid.uuid4().hex}_merged.pdf")

        writer = pypdf.PdfWriter()
        for path in input_paths:
            reader = pypdf.PdfReader(path)
            for page in reader.pages:
                writer.add_page(page)

        with open(output_path, "wb") as f:
            writer.write(f)

        return output_path
